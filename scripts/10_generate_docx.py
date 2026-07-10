"""
Renders manuscript_draft.md into a .docx (Word) document, mirroring
09_generate_pdf.py's structure and anchor-based figure/table insertion
points so both outputs stay in sync with the same source markdown.

Uses python-docx (pure Python, no native dependencies).
"""
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

PROJECT_DIR = Path(__file__).resolve().parent.parent
MD_PATH = PROJECT_DIR / "manuscript_draft.md"
FIGURES_DIR = PROJECT_DIR / "outputs" / "figures"
TABLES_DIR = PROJECT_DIR / "outputs" / "tables"
OUT_PATH = PROJECT_DIR / "VitalDB_Hypothermia_Manuscript.docx"

TEXT_GRAY = RGBColor(0x33, 0x33, 0x33)


def add_markdown_runs(paragraph, text):
    """Split **bold**/*italic*/`code` into runs within a single paragraph."""
    tokens = re.split(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            run = paragraph.add_run(tok[2:-2])
            run.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            run = paragraph.add_run(tok[1:-1])
            run.italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            run = paragraph.add_run(tok[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(tok)


def add_csv_table(doc, csv_path, max_cols=6):
    df = pd.read_csv(csv_path)
    if df.shape[1] > max_cols:
        df = df.iloc[:, :max_cols]
    table = doc.add_table(rows=1, cols=df.shape[1])
    table.style = "Light Grid Accent 1"
    for j, col in enumerate(df.columns):
        table.rows[0].cells[j].text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(7.5)
    doc.add_paragraph()


def add_figure(doc, fname, caption):
    doc.add_picture(str(FIGURES_DIR / fname), width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = TEXT_GRAY
    doc.add_paragraph()


ANCHORS = [
    ("2,567 cases were analyzed for Aims 1", "figure_table", [
        ("figure1_strobe.png", "Figure 1. STROBE cohort flow diagram: 6,388 total cases through "
         "candidate (N=2,824), analytic (N=2,568), and final Aim 1-3 (N=2,567) cohorts."),
        ("__table__", "table1_baseline_by_tertile.csv", "Table 1. Baseline characteristics by hypothermia-burden tertile."),
    ]),
    ("roughly a third higher odds of transfusion per clinically meaningful increase in exposure", "figure", [
        ("figure2_coefficient_plot.png", "Figure 2. Aim 1 primary model coefficients (SD-standardized), "
         "blood loss (linear) and transfusion (logistic)."),
    ]),
    ("a real but small deviation from linearity that adds texture to", "figure", [
        ("figure3_rcs_dose_response.png", "Figure 3. Restricted cubic spline: hypothermia burden vs. predicted blood loss."),
    ]),
    ("nearly half the variance in directly measured blood loss accounts for almost none", "figure_table", [
        ("__table__", "table9_v1_replication_comparison.csv", "Table 9. Delta Hb vs. EBL, identical covariates/cohort/model type."),
    ]),
]


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"

    text = MD_PATH.read_text(encoding="utf-8")
    text = text.split("## Notes for finalization")[0]
    lines = text.split("\n")

    title_done = False
    in_refs = False
    para_buffer = []

    def flush_para():
        if para_buffer:
            joined = " ".join(para_buffer).strip()
            para_buffer.clear()
            if not joined:
                return
            p = doc.add_paragraph()
            if in_refs:
                for r in p.runs:
                    r.font.size = Pt(9)
                p.paragraph_format.space_after = Pt(4)
            add_markdown_runs(p, joined)
            if in_refs:
                for r in p.runs:
                    r.font.size = Pt(9)

            for marker, kind, items in ANCHORS:
                if marker in joined:
                    for item in items:
                        if item[0] == "__table__":
                            _, csv_name, caption = item
                            cap_p = doc.add_paragraph()
                            cr = cap_p.add_run(caption)
                            cr.bold = True
                            cr.font.size = Pt(9)
                            add_csv_table(doc, TABLES_DIR / csv_name)
                        else:
                            fname, caption = item
                            add_figure(doc, fname, caption)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not title_done and line.startswith("# "):
            h = doc.add_heading(line[2:], level=0)
            i += 1
            if i < len(lines) and lines[i].strip().startswith("*"):
                sub = doc.add_paragraph()
                sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = sub.add_run(lines[i].strip().strip("*"))
                r.italic = True
                i += 1
            title_done = True
            continue

        if line.strip() == "---":
            flush_para()
            i += 1
            continue

        if line.startswith("## "):
            flush_para()
            heading = line[3:].strip()
            in_refs = heading.lower().startswith("references")
            if in_refs:
                doc.add_page_break()
            doc.add_heading(heading, level=1)
            i += 1
            continue

        if line.startswith("### "):
            flush_para()
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        if line.strip().startswith("- "):
            flush_para()
            p = doc.add_paragraph(style="List Bullet")
            add_markdown_runs(p, line.strip()[2:])
            i += 1
            continue

        if in_refs and re.match(r"^\d+\.\s", line.strip()):
            flush_para()
            p = doc.add_paragraph()
            add_markdown_runs(p, line.strip())
            for r in p.runs:
                r.font.size = Pt(9)
            i += 1
            continue

        if line.strip() == "":
            flush_para()
            i += 1
            continue

        para_buffer.append(line.strip())
        i += 1

    flush_para()
    doc.save(str(OUT_PATH))
    print(f"DOCX written -> {OUT_PATH} ({OUT_PATH.stat().st_size/1024:.0f} KB)")


if __name__ == "__main__":
    build()
